import numpy as np
import time
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])


# 模块五
def word_physics(save_path, **kwargs):
    """
    Physics-based Model word-report to create
    :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    # rul = 1
    # opt = 1

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of Physics-based Model", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Output data of Physics-based Model", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Output data", "Data shape"),
        ("Vibration signal", str(kwargs["all_data"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2
    document.add_heading("2. Parameters of the ball", level=1)
    # document.add_heading("1. Method of use", level=2)
    table_rows = [
        ("D", str(kwargs['D'])+" "+"mm"),
        ("Di", str(kwargs['Di'])+" "+"mm"),
        ("Do", str(kwargs['Do'])+" "+"mm"),
        ("Kb", str(kwargs['Kb'])+" "+"N·mm^−1"),
        ("Nb", str(kwargs['Nb'])),
        ("α", str(kwargs['α'])+" "+"deg"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格3
    document.add_heading("3. Parameters of the system", level=1)
    table_rows = [
        ("Ms", str(kwargs['Ms'])+" "+"kg"),
        ("Mp", str(kwargs['Mp'])+" "+"kg"),
        ("Mr", str(kwargs['Mr'])+" "+"kg"),
        ("Ks", str(kwargs['Ks'])+" "+"N·m^−1"),
        ("Kp", str(kwargs['Kp'])+" "+"N·m^−1"),
        ("Kr", str(kwargs['Kr'])+" "+"N·m^−1"),
        ("Cs", str(kwargs['Cs'])+" "+"Ns·m^−1"),
        ("Cp", str(kwargs['Cp'])+" "+"Ns·m^−1"),
        ("Cr", str(kwargs['Cr'])+" "+"Ns·m^−1"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格4.1
    document.add_heading("4. Parameters of defect definition", level=1)

    if kwargs['outer_ring_switch'] == 1:

        document.add_heading("Parameters to define the outer ring defect", level=2)
        table_rows = [
            ("Number", str(kwargs['outer_ring_number'])),
            ("Position", str(kwargs['outer_ring_local_position'])+" "+"deg"),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    # 表格4.2
    if kwargs['inner_ring_switch'] == 1:

        document.add_heading("Parameters to define the inner ring defect", level=2)
        table_rows = [
            ("Number", str(kwargs['inner_ring_number'])),
            ("Position", str(kwargs['inner_ring_local_position'])+" "+"deg"),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    # 表格4.3
    if kwargs['ball_switch'] == 1:

        document.add_heading("Parameters to define the ball defect", level=2)
        table_rows = [
            ("Number", str(kwargs['ball_number'])),
            ("Position", str(kwargs['ball_local_position'])+" "+"deg"),
            ("Identifier", str(kwargs['ball_fault_ball_identifier'])),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    # 表格4.4
    if kwargs['outer_ring_switch'] + kwargs['inner_ring_switch'] + kwargs['ball_switch'] > 0:

        document.add_heading("Parameters to define the defect size", level=2)
        table_rows = [
            ("L", str(kwargs['L'])+" "+"mm"),
            ("B", str(kwargs['B'])+" "+"mm"),
            ("H", str(kwargs['H'])+" "+"mm"),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    # 表格4.5
    if kwargs['segmentation_phy'] == 1:

        document.add_heading("Parameters to define the segmentation", level=2)
        table_rows = [
            ("Shift", str(kwargs['shift'])),
            ("Sample length", str(kwargs['length'])),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    # 表格4.6
    document.add_heading("Parameters to define the working condition", level=2)
    table_rows = [
        ("Fr", str(kwargs['Fr'])+" "+"N"),
        ("Fa", str(kwargs['Fa'])+" "+"N"),
        ("ω shaft", str(kwargs['omega_shaft'])+" "+"Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格4.7
    document.add_heading("Parameters to determine the simulation", level=2)
    table_rows = [
        ("Duration", str(kwargs['sim_duration'])+" "+"s"),
        ("Step size", str(kwargs['step_size'])+" "+"s"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格4.8
    document.add_heading("Parameters to 5-DoF bearing model", level=2)
    table_rows = [
        ("Mutation percentage", str(kwargs['mutation_percentage'])),
        ("Initial angular position", str(kwargs['initial_angular_position'])),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)


    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片：
    if kwargs['outer_ring_switch'] == 1:
        document.add_heading("Vibration signal and its envelope spectrum of outer defect", level=2)
        image_path = os.path.join(save_path, "time and frequency domain acceleration0.png")
        document.add_picture(image_path_or_stream=image_path, )

    if kwargs['inner_ring_switch'] == 1:
        document.add_heading("Vibration signal and its envelope spectrum of inner defect", level=2)
        image_path = os.path.join(save_path, "time and frequency domain acceleration1.png")
        document.add_picture(image_path_or_stream=image_path, )

    if kwargs['ball_switch'] == 1:
        document.add_heading("Vibration signal and its envelope spectrum of ball defect", level=2)
        image_path = os.path.join(save_path, "time and frequency domain acceleration2.png")
        document.add_picture(image_path_or_stream=image_path, )

    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if kwargs['output_file'] == 0:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + '.mat', "Data of physical-based model"),
            ("physical_based_bearing_defect_model_label" + '.mat', "Label of data"),
        ]
    elif kwargs['output_file'] == 1:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + '.xlsx', "Data of physical-based model"),
            ("physical_based_bearing_defect_model_label" + '.xlsx', "Label of data"),
        ]
    elif kwargs['output_file'] == 2:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + '.npy', "Data of physical-based model"),
            ("physical_based_bearing_defect_model_label" + '.npy', "Label of data"),
        ]
    elif kwargs['output_file'] == 3:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".csv", "Data of physical-based model"),
            ("physical_based_bearing_defect_model_label" + ".csv", "Label of data"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".txt", "Data of physical-based model"),
            ("physical_based_bearing_defect_model_label" + ".txt", "Label of data"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 图
    document.add_heading("Result image", level=2)
    if kwargs['output_image'] == 0:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".png",
             "Image files, 0-outer defect, 1-inner defect, 2-ball defect"),
        ]
    elif kwargs['output_image'] == 1:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".jpg",
             "Image files, 0-outer defect, 1-inner defect, 2-ball defect"),
        ]
    elif kwargs['output_image'] == 2:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".svg",
             "Image files, 0-outer defect, 1-inner defect, 2-ball defect"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("physical_based_bearing_defect_model" + ".pdf",
             "Image files, 0-outer defect, 1-inner defect, 2-ball defect"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of Physics-based Model.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_signal(save_path, **kwargs):
    """
    signal_based word-report to create
    :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    # rul = 1
    # opt = 1

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of Signal-based Model", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Output data of Signal-based Model", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Output data", "Data shape"),
        ("Vibration signal", str(kwargs["all_data"].shape)),
        ("Label", str(kwargs["all_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Parameters definition", level=1)

    document.add_heading("Bearing parameter", level=2)
    table_rows = [
        ("D", str(kwargs['D'])+" "+"mm"),
        ("Di", str(kwargs['di'])+" "+"mm"),
        ("Do", str(kwargs['do'])+" "+"mm"),
        ("Z", str(kwargs['Z'])),
        ("Type factor", str(kwargs['bearing_type_factor'])),
        ("α", str(kwargs['α'])+" "+"deg"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格2.2
    document.add_heading("Condition Parameter", level=2)
    table_rows = [
        ("Load max", str(kwargs['load_max'])+" "+"N"),
        ("Load proportional factor", str(kwargs['load_proportional_factor'])),
        ("Speed", str(kwargs['shaft_speed'])+" "+"Hz"),
        ("Resonance frequency", str(kwargs['resonance_frequency'])+" "+"Hz"),
        ("φ Limit", str(kwargs['phi_limit'])+" "+"deg"),
        ("Load distribution", str(kwargs['load_distribution_parameter'])),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格2.3
    document.add_heading("Defect Parameter", level=2)
    table_rows = [
        ("B", str(kwargs['B'])),
        ("Defect initial position", str(kwargs['defect_initial_position'])+" "+"deg"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)


    # 表格2.4
    document.add_heading("Simulation Parameter", level=2)
    table_rows = [
        ("Step size", str(kwargs['step_size'])+" "+"s"),
        ("Duration", str(kwargs['duration'])+" "+"s"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 表格2.5

    if kwargs['segmentation_sig'] == 1:
        document.add_heading("Segmentation parameter", level=2)
        table_rows = [
            ("Shift", str(kwargs['shift'])),
            ("Sample length", str(kwargs['length'])),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)



    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片：
    document.add_heading("Signal-based bearing vibration signal", level=2)

    if kwargs['outer_sig'] == 1:
        image_path = os.path.join(save_path, "signal_based_bearing_defect_model1.png")
        document.add_picture(image_path_or_stream=image_path, )

    if kwargs['inner_sig'] == 1:
        image_path = os.path.join(save_path, "signal_based_bearing_defect_model2.png")
        document.add_picture(image_path_or_stream=image_path, )

    if kwargs['ball_sig'] == 1:
        image_path = os.path.join(save_path, "signal_based_bearing_defect_model3.png")
        document.add_picture(image_path_or_stream=image_path, )


    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if kwargs['output_file'] == 0:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + '.mat', "Data of signal-based model"),
            ("signal_based_bearing_defect_model_label" + '.mat', "Label of data"),
        ]
    elif kwargs['output_file'] == 1:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + '.xlsx', "Data of signal-based model"),
            ("signal_based_bearing_defect_model_label" + '.xlsx', "Label of data"),
        ]
    elif kwargs['output_file'] == 2:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + '.npy', "Data of signal-based model"),
            ("signal_based_bearing_defect_model_label" + '.npy', "Label of data"),
        ]
    elif kwargs['output_file'] == 3:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".csv", "Data of signal-based model"),
            ("signal_based_bearing_defect_model_label" + ".csv", "Label of data"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".txt", "Data of signal-based model"),
            ("signal_based_bearing_defect_model_label" + ".txt", "Label of data"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if kwargs['output_image'] == 0:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".png", "Image files, 1-outer defect, 2-inner defect, 3-ball defect"),
        ]
    elif kwargs['output_image'] == 1:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".jpg", "Image files, 1-outer defect, 2-inner defect, 3-ball defect"),
        ]
    elif kwargs['output_image'] == 2:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".svg", "Image files, 1-outer defect, 2-inner defect, 3-ball defect"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("signal_based_bearing_defect_model" + ".pdf", "Image files, 1-outer defect, 2-inner defect, 3-ball defect"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of Signal-based Model.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


if __name__ == '__main__':
    arr = np.array([1, 2, 3, 4, 5, 6])


