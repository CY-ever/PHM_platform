import time
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])


def word_feature_selection(inputdata,outputdata,input_labels,output_labels,name_list, Features_selection,svd_dimension, pca_method,pca_dimension_method,pca_dimension,pca_percent,fda_dim,AE_encoding_dim,Monotonicity_threshold,Correlation_threshold,output_file,output_image,save_path):

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of feature selection", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for feature selection", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    if name_list is not None:
        name_content = list(name_list)
    else:
        name_content=None
    if input_labels is None:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_labels", str(None)),
            ("input_names", str(name_content)),
            ("output_data", str(outputdata.shape)),
            ("output_labels", str(None)),
        ]
    else:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_labels", str(input_labels.shape)),
            ("input_names", str(name_content)),
            ("output_data", str(outputdata.shape)),
            ("output_labels", str(output_labels.shape)),
        ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    i=3
    document.add_heading("2. Method information", level=1)
    if Features_selection==0:
        document.add_heading("Correlation", level=2)
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Threshold",str(Correlation_threshold)),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
        if name_list is not None:
            document.add_heading("3. Results", level=1)
            # 3.2. 放置结果图片
            document.add_heading("Correlation", level=2)
            image_path = os.path.join(save_path, "feature_selection.png")
            inline_shape = document.add_picture(image_path_or_stream=image_path)
            inline_shape.height = int(document.inline_shapes[0].height*0.25)
            inline_shape.width = int(document.inline_shapes[0].width*0.25)
            i=i+1
    elif Features_selection == 1:
        document.add_heading("Monotonicity", level=2)
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Threshold", str(Monotonicity_threshold)),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
        if name_list is not None:
            document.add_heading("3. Results", level=1)
            # 3.2. 放置结果图片
            document.add_heading("Monotonicity", level=2)
            image_path = os.path.join(save_path, "feature_selection.png")
            inline_shape = document.add_picture(image_path_or_stream=image_path)
            inline_shape.height = int(document.inline_shapes[0].height*0.25)
            inline_shape.width = int(document.inline_shapes[0].width*0.25)
            i = i + 1
    elif Features_selection == 2:
        document.add_heading("Singular value decomposition", level=2)
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Dimension", str(svd_dimension)),
        ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    elif Features_selection == 3:
        if pca_method==0:
            PCA_method_selection="EIG"
        else:
            PCA_method_selection = "SVD"

        document.add_heading("Principal component analysis", level=2)
        if pca_dimension_method == 0:
            table_rows = [
                ("Parameter", "Parameter value"),
                ("PCA method", PCA_method_selection),
                ("Dimension method", "Dimension"),
                ("Dimension",str(pca_dimension)),
            ]
        elif pca_dimension_method == 1:
            table_rows = [
                ("Parameter", "Parameter value"),
                ("PCA method", PCA_method_selection),
                ("Dimension method", "Percent"),
                ("Dimension", str(pca_percent)),
            ]
        else:
            table_rows = [
                ("Parameter", "Parameter value"),
                ("PCA method", PCA_method_selection),
                ("Dimension method", "Mle"),
            ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    elif Features_selection == 4:
        document.add_heading("Fisher discriminant analysis", level=2)
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Dimension", str(fda_dim)),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    else:
        document.add_heading("Autoencoder", level=2)
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Dimension", str(AE_encoding_dim)),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

    document.add_heading("%d. Save files"%i, level=1)
    document.add_heading("Result data", level=2)
    if Features_selection == 0:
        if output_file==0:
            if output_labels is None:

                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + ".mat", "Correlation coefficients of selected features"),
                ]

            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + ".mat", "Correlation coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".mat","The labels of data"),
                ]

        elif output_file==1:

            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + '.xlsx', "Correlation coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + '.xlsx', "Correlation coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".xlsx", "The labels of data"),
                ]

        elif output_file == 2:

            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation"+'.npy', "Correlation coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + '.npy', "Correlation coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".npy", "The labels of data"),
                ]

        elif output_file ==3:

            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation"+".csv", "Correlation coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + ".csv", "Correlation coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".npy", "The labels of data"),
                ]

        else:

            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation"+".txt", "Correlation coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Correlation" + ".txt", "Correlation coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".txt", "The labels of data"),
                ]

        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

        if name_list is not None:
            document.add_heading("Result image", level=2)
            if output_image==0:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".png", "Correlation histogram"),
                ]
            elif output_image==1:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".jpg", "Correlation histogram"),
                ]
            elif output_image==2:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".svg" "Correlation histogram"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".pdf", "Correlation histogram"),
                ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        else:
            pass
    elif Features_selection == 1:
        if output_file == 0:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".mat", "Monotonicity coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".mat", "Monotonicity coefficients of selected features"),
                    ("Labels_after_feature_selection" + ".mat", "The labels of data"),
                ]

        elif output_file == 1:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + '.xlsx', "Monotonicity coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + '.xlsx', "Monotonicity coefficients of selected features"),
                    ("Labels_after_feature_selection" + '.xlsx', "The labels of data"),
                ]
        elif output_file == 2:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + '.npy', "Monotonicity coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + '.npy', "Monotonicity coefficients of selected features"),
                    ("Labels_after_feature_selection" + '.npy', "The labels of data"),
                ]
        elif output_file == 3:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".csv", "Monotonicity coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".csv", "Monotonicity coefficients of selected features"),
                    ("Labels_after_feature_selection" + '.csv', "The labels of data"),
                ]
        else:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".txt", "Monotonicity coefficients of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Monotonicity" + ".txt", "Monotonicity coefficients of selected features"),
                    ("Labels_after_feature_selection" + '.txt', "The labels of data"),
                ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)

        if name_list is not None:
            document.add_heading("Result image", level=2)
            if output_image == 0:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".png", "Monotonicity histogram"),
                ]
            elif output_image == 1:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".jpg", "Monotonicity histogram"),
                ]
            elif output_image == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".svg" "Monotonicity histogram"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("feature_selection" + ".pdf", "Monotonicity histogram"),
                ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        else:
            pass
    elif Features_selection==2:
        if output_file == 0:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".mat", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".mat", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.mat', "The labels of data"),
                ]
        elif output_file == 1:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD"+ '.xlsx', "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + '.xlsx', "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.xlsx', "The labels of data"),
                ]
        elif output_file == 2:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + '.npy', "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + '.npy', "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.npy', "The labels of data"),
                ]
        elif output_file == 3:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".csv", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".csv", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.csv', "The labels of data"),
                ]
        else:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".txt", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("SVD" + ".txt", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.txt', "The labels of data"),
                ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    elif Features_selection==3:
        if output_file == 0:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".mat", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".mat", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.mat', "The labels of data"),
                ]
        elif output_file == 1:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".xlsx", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".xlsx", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.xlsx', "The labels of data"),
                ]
        elif output_file == 2:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".npy", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".npy", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.npy', "The labels of data"),
                ]
        elif output_file == 3:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".csv", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".csv", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.csv', "The labels of data"),
                ]
        else:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".txt", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("PCA" + ".txt", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.txt', "The labels of data"),
                ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    elif Features_selection==4:
        if output_file == 0:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".mat", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".mat", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.mat', "The labels of data"),
                ]
        elif output_file == 1:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".xlsx", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".xlsx", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.xlsx', "The labels of data"),
                ]
        elif output_file == 2:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".npy", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".npy", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.npy', "The labels of data"),
                ]
        elif output_file == 3:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".csv", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".csv", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.csv', "The labels of data"),
                ]
        else:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".txt", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("FDA" + ".txt", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.txt', "The labels of data"),
                ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)
    else:
        if output_file == 0:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".mat", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".mat", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.mat', "The labels of data"),
                ]
        elif output_file == 1:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".xlsx", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".xlsx", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.xlsx', "The labels of data"),
                ]
        elif output_file == 2:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".npy", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".npy", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.npy', "The labels of data"),
                ]
        elif output_file == 3:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".csv", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".csv", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.csv', "The labels of data"),
                ]
        else:
            if output_labels is None:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".txt", "Features after dimensionality reduction"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("AE" + ".txt", "Features after dimensionality reduction"),
                    ("Labels_after_feature_selection" + '.txt', "The labels of data"),
                ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table_full(table_rows, table)



    # 4. word文档保存
    file_name = "Report of feature selection.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)