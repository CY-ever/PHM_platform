import time
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])


def word_feature_extraction(inputdata,inputlabels,outputdata,output_labels, t_features_selected,f_features_selected,f_features_list,all_name,DWT_energy_p1,DWT_energy_p2,DWT_singular_p1,DWT_singular_p2,WP_energy_p1,WP_energy_p2,WP_singular_p1,WP_singular_p2,OPFCF_fault_type_list,OPFCF_fr, OPFCF_order, OPFCF_fs, OPFCF_switch, OPFCF_delta_f0, OPFCF_threshold, OPFCF_k, OPFCF_n_ball, OPFCF_d_ball, OPFCF_d_pitch, OPFCF_alpha,EMD_fr, EMD_n_ball, EMD_d_ball, EMD_d_pitch, EMD_alpha,EMD_fs,EMD_fault_type,EMD_n,EMD_ord,EMD_limit, FCF_nlevel, FCF_order, FCF_fs, FCF_fr, FCF_n_ball, FCF_d_ball, FCF_d_pitch, FCF_alpha,FCF_output_image,output_file,save_path):

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of feature extraction", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for feature extraction", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    if inputlabels is None:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_labels", str(None)),
            ("output_data", str(outputdata.shape)),
            ("output_labels", str(None)),
        ]
    else:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_labels", str(inputlabels.shape)),
            ("output_data", str(outputdata.shape)),
            ("output_labels", str(output_labels.shape)),
        ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Selected features", level=1)
    table_rows = [
        ("Feature names",),
        (str(all_name),),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    if f_features_selected==1:
        # 列出频域特征的具体参数配置
        i=0
        if f_features_list[0] == 1:
            i=i+1
            document.add_heading("%d. Discrete wavelet transform energy entropy"%i, level=2)
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Wavelet basis", str(DWT_energy_p1)),
                ("Level", str(DWT_energy_p2)),
            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        if f_features_list[1] == 1:
            i = i + 1
            document.add_heading("%d. Discrete wavelet transform singular entropy"%i, level=2)
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Wavelet basis", str(DWT_singular_p1)),
                ("Level", str(DWT_singular_p2)),
            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        if f_features_list[2] == 1:
            i = i + 1
            document.add_heading("%d. Wavelet packet energy entropy"%i, level=2)
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Wavelet basis", str(WP_energy_p1)),
                ("Level", str(WP_energy_p2)),
            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        if f_features_list[3] == 1:
            i = i + 1
            document.add_heading("%d. Wavelet packet singular entropy" % i, level=2)
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Wavelet basis", str(WP_singular_p1)),
                ("Level", str(WP_singular_p2)),
            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        if f_features_list[4] == 1:
            i = i + 1
            document.add_heading("%d. Occurrence probability of fault characteristic frequency" % i, level=2)
            if OPFCF_fault_type_list[0]==1:
                OPFCF_fault_type="BPFO"
            elif OPFCF_fault_type_list[1]==1:
                OPFCF_fault_type = "BPFI"
            elif OPFCF_fault_type_list[2] == 1:
                OPFCF_fault_type = "BSF"
            else:
                OPFCF_fault_type = "FTF"
            if OPFCF_switch==0:
                table_rows = [
                    ("Parameter", "Parameter value"),
                    ("Fault type", OPFCF_fault_type),
                    ("Order", str(OPFCF_order)),
                    ("Sampling frequency", str(OPFCF_fs)+ " " + "Hz"),
                    ("Rotation frequency", str(OPFCF_fr)+ " " + "Hz"),
                    ("Interval", "Frequency value interval"),
                    ("Fixed frequency interval", str(OPFCF_delta_f0)+ " " + "Hz"),
                    ("Threshold", str(OPFCF_threshold)),
                    ("Number of rolling elements",str(OPFCF_n_ball)),
                    ("Ball diameter", str(OPFCF_d_ball)+ " " + "mm"),
                    ("Pitch diameter", str(OPFCF_d_pitch)+ " " + "mm"),
                    ("Initial contact angle", str(OPFCF_alpha)+ " " + "°"),
                ]
                table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table_full(table_rows, table)
            else:
                table_rows = [
                    ("Parameter", "Parameter value"),
                    ("Fault type", OPFCF_fault_type),
                    ("Order",  str(OPFCF_order)),
                    ("Sampling frequency", str(OPFCF_fs)+ " " + "Hz"),
                    ("Rotation frequency", str(OPFCF_fr)+ " " + "Hz"),
                    ("Interval", "Percentage interval"),
                    ("Threshold", str(OPFCF_threshold)),
                    ("Percent", str(OPFCF_k)),
                    ("Number of rolling elements", str(OPFCF_n_ball)),
                    ("Ball diameter", str(OPFCF_d_ball)+ " " + "mm"),
                    ("Pitch diameter",str(OPFCF_d_pitch)+ " " + "mm"),
                    ("Initial contact angle", str(OPFCF_alpha)+ " " + "°"),
                ]
                table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table_full(table_rows, table)
        if f_features_list[7] == 1:
            i = i + 1
            document.add_heading("%d. Envelope spectrum" % i, level=2)
            if EMD_fault_type==0:
                EMD_fault="BPFO"
            elif EMD_fault_type==1:
                EMD_fault = "BPFI"
            elif EMD_fault_type==2:
                EMD_fault = "BSF"
            else:
                EMD_fault = "FTF"
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Rotation frequency", str(EMD_fr)+ " " + "Hz"),
                ("Number of rolling elements", str(EMD_n_ball)),
                ("Ball diameter", str(EMD_d_ball)+ " " + "mm"),
                ("Pitch diameter", str(EMD_d_pitch)+ " " + "mm"),
                ("Initial contact angle", str(EMD_alpha)+ " " + "°"),
                ("Sampling frequency", str(EMD_fs)+ " " + "Hz"),
                ("Fault type", EMD_fault),
                ("Range of peak detection", str(EMD_n)),
                ("Order", str(EMD_ord)),
                ("Limit", str(EMD_limit)+ " " + "Hz"),
            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)
        if f_features_list[8] == 1:
            i = i + 1
            document.add_heading("%d. Fault characteristic frequency ratio " % i, level=2)
            table_rows = [
                ("Parameter", "Parameter value"),
                ("Level", str(FCF_nlevel)),
                ("Order", str(FCF_order)),
                ("Sampling frequency", str(FCF_fs)+ " " + "Hz"),
                ("Rotation frequency", str(FCF_fr)+ " " + "Hz"),
                ("Number of rolling elements", str(FCF_n_ball)),
                ("Ball diameter", str(FCF_d_ball)+ " " + "mm"),
                ("Pitch diameter", str(FCF_d_pitch)+ " " + "mm"),
                ("Initial contact angle", str(FCF_alpha)+ " " + "°"),
                ("FCF-ratio Image", str(FCF_output_image)),

            ]
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table_full(table_rows, table)

    document.add_heading("3. Save files", level=1)
    document.add_heading("Result data", level=2)

    if t_features_selected == 1 and f_features_selected == 1:
        if output_labels is None:
            if output_file == 0:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features" + ".mat", "Selected features"),
                    ("All_feature_names" + ".mat", "The names of selected features"),
                ]
            elif output_file == 1:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features" + '.xlsx', "Selected features"),
                    ("All_feature_names" + '.xlsx', "The names of selected features"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features" + '.npy', "Selected features"),
                    ("All_feature_names" + '.npy', "The names of selected features"),
                ]
            elif output_file == 3:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features" + ".csv", "Selected features"),
                    ("All_feature_names" + ".csv", "The names of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features" + ".txt", "Selected features"),
                    ("All_feature_names" + ".txt", "The names of selected features"),
                ]
        else:
            if output_file==0:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features"+".mat", "Selected features"),
                    ("All_feature_names"+".mat","The names of selected features"),
                    ("Labels_after_feature_extraction"+".mat","The labels of data"),
                ]
            elif output_file==1:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features"+'.xlsx', "Selected features"),
                    ("All_feature_names"+'.xlsx', "The names of selected features"),
                    ("Labels_after_feature_extraction" + '.xlsx', "The labels of data"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features"+'.npy', "Selected features"),
                    ("All_feature_names"+'.npy', "The names of selected features"),
                    ("Labels_after_feature_extraction" + '.npy', "The labels of data"),
                ]
            elif output_file ==3:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features"+".csv", "Selected features"),
                    ("All_feature_names"+".csv", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".csv", "The labels of data"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("All_features"+".txt", "Selected features"),
                    ("All_feature_names"+".txt", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".txt", "The labels of data"),
                ]
    elif t_features_selected == 1 and f_features_selected == 0:
        if output_labels is None:
            if output_file == 0:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features" + ".mat", "Selected features"),
                    ("Time_domain_feature_names" + ".mat", "The names of selected features"),
                ]
            elif output_file == 1:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features" + '.xlsx', "Selected features"),
                    ("Time_domain_feature_names" + '.xlsx', "SThe names of selected features"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features" + '.npy', "Selected features"),
                    ("Time_domain_feature_names" + '.npy', "The names of selected features"),
                ]
            elif output_file == 3:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features" + ".csv", "Selected features"),
                    ("Time_domain_feature_names" + ".csv", "The names of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features" + ".txt", "Selected features"),
                    ("Time_domain_feature_names" + ".txt", "The names of selected features"),
                ]
        else:
            if output_file==0:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features"+".mat", "Selected features"),
                    ("Time_domain_feature_names"+".mat","The names of selected features"),
                    ("Labels_after_feature_extraction" + ".mat", "The labels of data"),
                ]
            elif output_file==1:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features"+'.xlsx', "Selected features"),
                    ("Time_domain_feature_names"+'.xlsx', "SThe names of selected features"),
                    ("Labels_after_feature_extraction" + '.xlsx', "The labels of data"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features"+'.npy', "Selected features"),
                    ("Time_domain_feature_names"+'.npy', "The names of selected features"),
                    ("Labels_after_feature_extraction" + '.npy', "The labels of data"),
                ]
            elif output_file ==3:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features"+".csv", "Selected features"),
                    ("Time_domain_feature_names"+".csv", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".csv", "The labels of data"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Time_domain_features"+".txt", "Selected features"),
                    ("Time_domain_feature_names"+".txt", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".txt","The labels of data"),
                ]
    elif t_features_selected == 0 and f_features_selected==1:
        if output_labels is None:
            if output_file == 0:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".mat", "Selected features"),
                    ("Frequency_feature_names" + ".mat", "The names of selected features"),
                ]
            elif output_file == 1:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + '.xlsx', "Selected features"),
                    ("Frequency_feature_names" + '.xlsx', "The names of selected features"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + '.npy', "Selected features"),
                    ("Frequency_feature_names" + '.npy', "The names of selected features"),
                ]
            elif output_file == 3:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".csv", "Selected features"),
                    ("Frequency_feature_names" + ".csv", "The names of selected features"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".txt", "Selected features"),
                    ("Frequency_feature_names" + ".txt", "The names of selected features"),
                ]
        else:
            if output_file == 0:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".mat", "Selected features"),
                    ("Frequency_feature_names" + ".mat", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".mat", "The labels of data"),
                ]
            elif output_file == 1:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + '.xlsx', "Selected features"),
                    ("Frequency_feature_names" + '.xlsx', "The names of selected features"),
                    ("Labels_after_feature_extraction" + '.xlsx',"The labels of data"),
                ]
            elif output_file == 2:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + '.npy', "Selected features"),
                    ("Frequency_feature_names" + '.npy', "The names of selected features"),
                    ("Labels_after_feature_extraction" + '.npy', "The labels of data"),
                ]
            elif output_file == 3:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".csv", "Selected features"),
                    ("Frequency_feature_names" + ".csv", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".csv", "The labels of data"),
                ]
            else:
                table_rows = [
                    ("Filename", "Description"),
                    ("Frequency_features" + ".txt", "Selected features"),
                    ("Frequency_feature_names" + ".txt", "The names of selected features"),
                    ("Labels_after_feature_extraction" + ".txt", "The labels of data"),
                ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)


    # 4. word文档保存
    file_name = "Report of feature extraction.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)