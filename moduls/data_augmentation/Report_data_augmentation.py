import numpy as np
import time
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from moduls.data_augmentation.writein import writein
# from JY_code_all.writein import writein
# from utils.read_data import r


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])

def word_GAN(data,label,newdata_all,newlabel_all,num, Z_dim,save_path,output_file,output_image):

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of generative adversarial networks (GAN)", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for GAN", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    try:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(data.shape)),
            ("input_label", str(label.shape)),
            ("output_data", str(newdata_all.shape)),
            ("output_label", str(newlabel_all.shape)),
        ]
    except:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(data.shape)),
            ("input_label", 'None'),
            ("output_data", str(newdata_all.shape)),
            ("output_label", 'None'),
        ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Num", str(num)),
        ("Input noise size", str(Z_dim)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Augmentation signal", level=2)
    image_path = os.path.join(save_path, "GAN_image2.png")
    inline_shape=document.add_picture(image_path_or_stream=image_path)
    inline_shape.height= int(document.inline_shapes[0].height * 0.50)
    inline_shape.width = int(document.inline_shapes[0].width * 0.50)

    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_data" + '.mat', "Augmentation signal after GAN"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_data" + '.xlsx', "Augmentation signal after GAN"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_data" + '.npy', "Augmentation signal after GAN"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_data" + ".csv", "Augmentation signal after GAN"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_data" + ".txt", "Augmentation signal after GAN"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_image2" + ".png",
             "Time-domain signal diagram after GAN"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_image2" + ".jpg",
             "Time-domain signal diagram after GAN"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_image2" + ".svg",
             "Time-domain signal diagram after GAN"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("GAN_image2" + ".pdf",
             "Time-domain signal diagram after GAN"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of GAN.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)

def word_image_transformation(inputdata,label,multi,newdata_all,newlabel_all,deltax, deltay,rot, snr,rescale,save_path,output_file,output_image):

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of image-transformation-based methods", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for image-transformation-based methods", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    try:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_label", str(label.shape)),
            ("output_data", str(newdata_all.shape)),
            ("output_label", str(newlabel_all.shape)),
        ]
    except:
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(inputdata.shape)),
            ("input_label", 'None'),
            ("output_data", str(newdata_all.shape)),
            ("output_label", 'None'),
        ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Translation", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Translation in X", str(deltax)),
        ("Translation in Y", str(deltay)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Rotation", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Rotation angle", str(rot)+ " " + "deg"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    document.add_heading("3. Noise", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Signal-to-noise ratio", str(snr)+ " " + "dB"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("4. Scale", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Scale factor", str(rescale)),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)


    document.add_heading("5. Number of augmentation", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Num", str(multi)),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)


    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Augmentation signal", level=2)
    image_path = os.path.join(save_path, "transformation_DA_image_new2.png")
    inline_shape=document.add_picture(image_path_or_stream=image_path)
    inline_shape.height= int(document.inline_shapes[0].height * 0.50)
    inline_shape.width = int(document.inline_shapes[0].width * 0.50)

    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Image_transformation_data" + '.mat', "Augmentation signal after image-transformation-based methods"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Image_transformation_data" + '.xlsx', "Augmentation signal after image-transformation-based methods"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Image_transformation_data" + '.npy', "Augmentation signal after image-transformation-based methods"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("Image_transformation_data" + ".csv", "Augmentation signal after image-transformation-based methods"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Image_transformation_data" + ".txt", "Augmentation signal after image-transformation-based methods"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("transformation_DA_image_new2" + ".png", "Time-domain signal diagram after image-transformation-based methods"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("transformation_DA_image_new2" + ".jpg", "Time-domain signal diagram after image-transformation-based methods"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("transformation_DA_image_new2" + ".svg", "Time-domain signal diagram after image-transformation-based methods"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("transformation_DA_image_new2" + ".pdf", "Time-domain signal diagram after image-transformation-based methods"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of Image_transformation.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)

def word_Monte_Carlo(data, output,mode,distribution,function_select,m,a, b,c,d,e,save_path,output_file,output_image):

    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of Monte Carlo sampling", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for Monte Carlo sampling", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    output=np.array(output)
    num_output=len(output)
    if mode==0:
        files = os.listdir(data)  # 读入文件夹
        num_files = len(files)

        table_rows = [
            ("Data", "Data description"),
            ("input_data", "%d degradation curves"% num_files),
            ("output_data", "%d degradation curves"% num_output),
        ]
    else :
        data=writein(data, 1)
        # data = np.array(data)
        table_rows = [
            ("Data", "Data shape"),
            ("input_data", str(data.shape)),
            ("output_data", str(output.shape)),
        ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    if mode==0:
        mode_name="Data fitting and Monte Carlo sampling"
    else:
        mode_name="Monte Carlo sampling"

    if distribution==0:
        distribution_name="Weibull distribution"
    elif distribution==1:
        distribution_name="Normal distribution"
    else:
        distribution_name="Gamma distribution"

    if function_select==0:
        function_name="a * exp(b * t) - 1"
    elif function_select==1:
        function_name = "a * exp(b * t) + c"
    elif function_select==2:
        function_name = "a * exp(b * t) + c * exp(d * t) + e"
    elif function_select==3:
        function_name = "a * t ^ 2 + b * t + c"
    elif function_select==4:
        function_name = "a * t ^ 3 + b * t ^ 2 + c * t + d"
    else:
        function_name = "a * exp(b * t) + c * t ^ 2 + d"

    document.add_heading("2. Method information", level=1)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Mode", mode_name),
        ("Distribution", distribution_name),
        ("Function select", function_name),
        ("Num", str(m)),
        ("Parameter a", str(a)),
        ("Parameter b", str(b)),
        ("Parameter c", str(c)),
        ("Parameter d", str(d)),
        ("Parameter e", str(e)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    if mode == 0:
        # 子标题三：结果信息；（数据，图，表）
        document.add_heading("3. Results", level=1)

        # 3.2. 放置结果图片
        document.add_heading("Original degradation trajectories", level=2)
        image_path1 = os.path.join(save_path, "Original_degradation_trajectories.png")
        inline_shape=document.add_picture(image_path_or_stream=image_path1)
        inline_shape.height= int(document.inline_shapes[0].height*0.7)
        inline_shape.width = int(document.inline_shapes[0].width*0.7)

        document.add_heading("Degradation trajectories after fitting", level=2)
        image_path2 = os.path.join(save_path, "Degradation_trajectories_after_fitting.png")
        inline_shape = document.add_picture(image_path_or_stream=image_path2)
        inline_shape.height = int(document.inline_shapes[0].height*0.9)
        inline_shape.width = int(document.inline_shapes[0].width*0.9)

        document.add_heading("Degradation trajectories after Monte Carlo sampling", level=2)
        image_path3 = os.path.join(save_path, "Degradation_trajectories_after_Monte_Carlo_sampling.png")
        inline_shape = document.add_picture(image_path_or_stream=image_path3)
        inline_shape.height = int(document.inline_shapes[0].height *0.9)
        inline_shape.width = int(document.inline_shapes[0].width*0.9 )

        # 子标题四：结果保存路径信息；（数据，图，表）
        # 数据
        document.add_heading("4. Save files", level=1)
        document.add_heading("Result data", level=2)
        if output_file == 0:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_data" + '.mat', "Augmentation signal after Monte Carlo sampling"),
            ]
        elif output_file == 1:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_data" + '.xlsx', "Augmentation signal after Monte Carlo sampling"),
            ]
        elif output_file == 2:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_data" + '.npy', "Augmentation signal after Monte Carlo sampling"),
            ]
        elif output_file == 3:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_data" + ".csv", "Augmentation signal after Monte Carlo sampling"),
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_data" + ".txt", "Augmentation signal after Monte Carlo sampling"),
            ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
        # 图

        document.add_heading("Result image", level=2)
        if output_image == 0:
            table_rows = [
                ("Filename", "Description"),
                ("Original_degradation_trajectories" + ".png",
                 "Original degradation trajectories"),
                ("Degradation_trajectories_after_fitting" + ".png",
                 "Degradation trajectories after fitting"),
                ("Degradation_trajectories_after_Monte_Carlo_sampling" + ".png",
                 "Degradation trajectories after Monte Carlo sampling"),
            ]
        elif output_image == 1:
            table_rows = [
                ("Filename", "Description"),
                ("Original_degradation_trajectories" + ".jpg",
                 "Original degradation trajectories"),
                ("Degradation_trajectories_after_fitting" + ".jpg",
                 "Degradation trajectories after fitting"),
                ("Degradation_trajectories_after_Monte_Carlo_sampling" + ".jpg",
                 "Degradation trajectories after Monte Carlo sampling"),
            ]
        elif output_image == 2:
            table_rows = [
                ("Filename", "Description"),
                ("Original_degradation_trajectories" + ".svg",
                 "Original degradation trajectories"),
                ("Degradation_trajectories_after_fitting" + ".svg",
                 "Degradation trajectories after fitting"),
                ("Degradation_trajectories_after_Monte_Carlo_sampling" + ".svg",
                 "Degradation trajectories after Monte Carlo sampling"),
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Original_degradation_trajectories" + ".pdf",
                 "Original degradation trajectories"),
                ("Degradation_trajectories_after_fitting" + ".pdf",
                 "Degradation trajectories after fitting"),
                ("Degradation_trajectories_after_Monte_Carlo_sampling" + ".pdf",
                 "Degradation trajectories after Monte Carlo sampling"),
            ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        # 子标题四：结果保存路径信息；（数据，图，表）
        # 数据
        document.add_heading("3. Save files", level=1)
        document.add_heading("Result data", level=2)
        if output_file == 0:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_parameters" + '.mat', "Monte Carlo sampling of input parameters "),
            ]
        elif output_file == 1:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_parameters" + '.xlsx', "Monte Carlo sampling of input parameters "),
            ]
        elif output_file == 2:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_parameters" + '.npy', "Monte Carlo sampling of input parameters "),
            ]
        elif output_file == 3:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_parameters" + ".csv", "Monte Carlo sampling of input parameters "),
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Monte_carlo_parameters" + ".txt", "Monte Carlo sampling of input parameters "),
            ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of Monte Carlo sampling.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


