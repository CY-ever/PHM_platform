import numpy as np
import time
import scipy.io as scio
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from moduls.ml.dataset import import_data

# from dataset import import_data
# from utils import to_cat

def file_exten(output_files):
    if output_files == 1:
        file_extension = ".xlsx"
    elif output_files == 2:
        file_extension = ".npy"
    elif output_files == 3:
        file_extension = ".csv"
    elif output_files == 4:
        image_extension = ".txt"
    elif output_files == 0:
        file_extension = ".mat"
    else:
        file_extension = ""

    return file_extension


def image_exten(output_images):
    if output_images == 1:
        image_extension = ".jpg"
    elif output_images == 2:
        image_extension = ".svg"
    elif output_images == 3:
        image_extension = ".pdf"
    elif output_images == 0:
        image_extension = ".png"
    else:
        image_extension = ""
    return image_extension

def part_4(document, kwargs):
    document.add_heading("4. Save files", level=1)
    # 4.1 数据文件
    file_extension = file_exten(kwargs['output_files'])
    document.add_heading("Result data", level=2)
    table_rows = [
        ("Filename", "Description"),
        ("traindata"+file_extension, "Training set of the model"),
        ("trainlabel"+file_extension, "Training labels of the model"),
        ("testdata"+file_extension, "Test set of the model"),
        ("testlabel"+file_extension, "Test labels of the model"),
        ("pred_label"+file_extension, "Predicted values for the test set"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    # 4.1 图片文件
    document.add_heading("Result image", level=2)
    image_extension = image_exten(kwargs['output_images'])
    table_rows = [
        ("Filename", "Description"),
    ]
    print(table_rows)
    if kwargs['opt_algorithm']:
        table_rows.append(("optimization_curve" + image_extension, "Optimisation curves"))
        print(table_rows)

    if kwargs['rul_pre']:
        table_rows.append(("rul_figure" + image_extension, "RUL prediction curve"))
        print(table_rows)

    else:
        table_rows.append(("confusion_matrix" + image_extension, "Confusion matrix"))
        print(table_rows)

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table_full(table_rows, table)

    return document

def dataset_test():
    file = "py/Dataset/data from Chaoren.mat"
    x_train_dict = scio.loadmat(file)
    x_train_all = x_train_dict.get('train_data')
    print()
    # y_train = to_cat(y_train_all, num_classes=4)

    # 样本数据形状
    data_path = './Dataset/'
    x_train, x_test, y_train, y_test = import_data(data_path, model='KNN')
    print("0:", np.sum(y_train == 0) + np.sum(y_test == 0))
    print("1:", np.sum(y_train == 1) + np.sum(y_test == 1))
    print("2:", np.sum(y_train == 2) + np.sum(y_test == 2))
    print("3:", np.sum(y_train == 3) + np.sum(y_test == 3))

    list_test = [[0], [1], [2], [3]]
    list_test = np.array(list_test)


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 2, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 2, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])

# 模块一
def word_signal():
    """
    signal_based word-report to create
    :return:
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report", level=0)

    # 2.1 生成时间（文本段落）
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一
    document.add_heading("1. Parameters of bearing defect simulation", level=1)
    document.add_heading("1. Bearing parameter structure", level=2)

    # 4. 创建表格
    # 4.1 表格数据准备
    table_rows = [
        ("D(mm)", "Z", "Di(mm)", "Do(mm)", "α(deg)", "Type factor"),
        ("22.225", "16", "102.7938", "147.7264", "40", "1.5")
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 表格二：
    document.add_heading("2. Condition parameter structure", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("Load max(N)", "Shaft speed(Hz)", "Load proportional factor"),
        ("1000", "25", "0.1"),
        ("Resonance frequency(Hz)", "Load distribution parameter", "Phi limit(deg)"),
        ("3000", "0.42", "82"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格三：
    document.add_heading("3. Defect parameter structure", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("Defect type", "B", "Defect initial position(deg)"),
        ("1", "300", "15"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格四：
    document.add_heading("4. Simlation parameter structure", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("Step size(s)", "Duration(s)"),
        ("0.0001", "1"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格五：
    document.add_heading("5. Defect frequency", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("Defect frequency",),
        ("172.816",),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 5. 添加图片
    document.add_heading("2. Result", level=1)
    document.add_picture(image_path_or_stream="./test/confusion_matrix.png",
                         )

    # 6. word文档保存
    document.save("./Report of signal-based model.doc")


def word_physics():
    """
    signal_based word-report to create
    :return:
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report", level=0)

    # 2.1 生成时间（文本段落）
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一
    document.add_heading("1. Parameters of bearing defect simulation", level=1)
    document.add_heading("1. Bearing parameter structure", level=2)

    # 4. 创建表格
    # 4.1 表格数据准备
    # TODO：在此处填入参数
    table_rows = [
        ("D(mm)", "Di(mm)", "Do(mm)", "Kb", "Nb", "α(deg)"),
        ("22.225", "16", "102.7938", "147.7264", "40", "1.5")
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 表格二：
    # TODO：在此处填入参数
    document.add_heading("2. System Parameter structure", level=2)
    table_rows = [

        ("Ms(kg)", "Mp(kg)", "Mr(kg)"),
        ("1000", "25", "0.1"),
        ("Ks(N/m)", "Kp(N/m)", "Kr(N/m)"),
        ("3000", "0.42", "82"),
        ("Cs(Ns/m)", "Cp(Ns/m)", "Cr(Ns/m)"),
        ("3000", "0.42", "82"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格三：
    document.add_heading("3. Defect parameter structure", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("L(mm)", "B(mm)", "H(mm)"),
        ("1", "300", "15"),
        ("Outer ring switch", "Inner ring switch", "Ball ring switch"),
        ("1", "300", "15"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格四：
    document.add_heading("4. Simlation parameter structure", level=2)
    table_rows = [
        # ("Condition parameter structure"),
        ("Step size(s)", "Duration(s)"),
        ("0.0001", "1"),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 5. 添加图片
    document.add_heading("2. Result", level=1)
    document.add_picture(image_path_or_stream="./test/confusion_matrix.png",
                         )

    # 6. word文档保存
    document.save("./Report of physics-based model.doc")


# 模块五
def word_knn(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
	○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
	○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning K-Nearest Neighbor", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for KNN", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("KNN", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of KNN", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("K", str(kwargs['K'])),
        ("Distance", str(kwargs['weights'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵
    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # document = part_4(document, kwargs=kwargs)

    # 6. word文档保存
    file_name = "Report of Machine_Learning_KNN.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)



def word_svm(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning svm word-report to create
        :return:
        """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Support Vector Machine", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for SVM", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("SVM", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of SVM", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("C", str(kwargs['C'])),
        ("gamma", str(kwargs['gamma'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter Value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_SVM.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_dt(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Decision Tree word-report to create
        :return:
        """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Decision Tree", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for Decision Tree", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("Decision Tree", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of Decision Tree", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Max depth", str(kwargs['max_depth'])),
        ("Max leaf nodes", str(kwargs['max_leaf_nodes'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_Decision_Tree.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_rf(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Random Forest word-report to create
        :return:
        """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Random Forest", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for Random Forest", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("Random Forest", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of Random Forest", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Max depth", str(kwargs['max_depth'])),
        ("Max leaf nodes", str(kwargs['max_leaf_nodes'])),
        ("Number of estimators", str(kwargs['n_estimators'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_Random_Forest.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_bagging(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Bagging word-report to create
        :return:
        """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Bagging", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for Bagging", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("Bagging", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of Bagging", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        # ("Max depth", str(kwargs['max_depth'])),
        ("Max leaf nodes", str(kwargs['max_leaf_nodes'])),
        ("Number of estimators", str(kwargs['n_estimators'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_Bagging.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_et(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Extra Tree word-report to create
        :return:
        """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Extra Tree", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for Extra Tree", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("Extra Tree", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of Extra Tree", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Max depth", str(kwargs['max_depth'])),
        ("Max leaf nodes", str(kwargs['max_leaf_nodes'])),
        ("Number of estimators", str(kwargs['n_estimators'])),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_Extra_Tree.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_ae(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Autoencoder word-report to create
        :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning Autoencoder", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for Autoencoder", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("Autoencoder", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of Autoencoder", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Layer count", str(kwargs['LayerCount'])),
        ("Epochs", str(kwargs['epochs'])),
        ("Batchsize", str(kwargs['batchSize'])),
        ("Dense activation", str(kwargs['denseActivation'])),
        # ("Max leaf nodes", str(kwargs['optimizer'])),
        # ("Number of estimators", str(kwargs['loss'])),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_Autoencoder.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_dbn(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Deep Belief Network word-report to create
        :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning DBN", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for DBN", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("DBN", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of DBN", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Dropout", str(kwargs['Dropout'])),
        ("Learning rate", str(kwargs['LearningRate_nn'])),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_DBN.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_cnn(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Convolutional Neural Network word-report to create
        :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning CNN", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for CNN", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("CNN", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of CNN", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("Dropout", str(kwargs['dropout'])),
        ("Learning rate:", str(kwargs['learning_rate'])),
        ("Batch size", str(kwargs['batch_size'])),
        ("Epochs", str(kwargs['epochs'])),
    ]


    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_CNN.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_lstm(save_path, **kwargs):
    """
    machine learning svm word-report to create
    :return:
    """
    """
        machine learning Long Short-Term Memory word-report to create
        :return:
    """
    # 这两个是判断参数：用于控制report是否生成某一表格或内容。可自行修改配置和删除。
    # 这里仅作说明
    """
    ○ 运行日志信息（author, date…)
    ○ 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    ○ 方法的信息：（方法类型，参数配置，）  ；
    ○ 结果信息；（数据，图，表）
    """
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of machine learning LSTM", level=0)

    # 2.1 运行日志信息
    print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"create datum: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)
    """
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input data for LSTM", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！
    table_rows = [
        ("Input data", "Data shape"),
        ("train_data", str(kwargs["train_data"].shape)),
        ("train_label", str(kwargs["train_label"].shape)),
        ("test_data", str(kwargs["test_data"].shape)),
        ("test_label", str(kwargs["test_label"].shape)),
    ]
    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Method of use", level=2)
    if kwargs["opt_algorithm"] == 1:
        opt_status = "PSO"
    elif kwargs["opt_algorithm"] == 2:
        opt_status = "GA"
    elif kwargs["opt_algorithm"] == 3:
        opt_status = "SA"
    else:
        opt_status = "None"

    table_rows = [
        ("Method", "Status values"),
        ("LSTM", "Used"),
        ("Optimizer", opt_status),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.2：
    document.add_heading("2. Parameters of LSTM", level=2)
    table_rows = [
        ("Parameters", "Parameter value"),
        ("LSTM count", str(kwargs['lstm_count'])),
        ("Dropout rate", str(kwargs['dropoutRate'])),
        ("Epochs", str(kwargs['epochs'])),
        ("Batch size", str(kwargs['batchSize'])),
        ("Dense activation", str(kwargs['denseActivation'])),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 表格2.3：
    if kwargs['opt_algorithm'] == 1:
        document.add_heading("3. Parameters of PSO optimization algorithm", level=2)
        table_rows = [
            ("Parameters of PSO", "Parameter value"),
            ("Pop size", str(kwargs['pso_pop_size'])),
            ("Max iteration", str(kwargs['pso_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)
    elif kwargs['opt_algorithm'] == 2:
        document.add_heading("3. Parameters of GA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of GA", "Parameter value"),
            ("Pop size", str(kwargs['ga_pop_size'])),
            ("Max iteration", str(kwargs['ga_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    elif kwargs['opt_algorithm'] == 3:
        document.add_heading("3. Parameters of SA optimization algorithm", level=2)
        table_rows = [
            ("Parameters of SA", "Parameter value"),
            ("Alpha", str(kwargs['sa_alpha'])),
            ("Max iteration", str(kwargs['sa_max_itr'])),
        ]
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

        table_full(table_rows, table)

    else:
        pass

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)
    # 表3.1：训练结果
    document.add_heading("Test accuracy", level=2)
    table_rows = [
        ("Accuracy", str(round(kwargs['accuracy'], 3))),
    ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 3.2. 放置结果图片：混淆矩阵

    # 3.3. 放置结果图片：优化曲线
    if kwargs['opt_algorithm']:
        img_path = os.path.join(save_path, "optimization_curve.png")
        document.add_heading("Optimization curve", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )

    if kwargs['rul_pre']:
        img_path = os.path.join(save_path, "rul_figure.png")
        document.add_heading("RUL prediction chart", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    else:
        img_path = os.path.join(save_path, "confusion_matrix.png")
        document.add_heading("Confusion matrix", level=2)
        document.add_picture(image_path_or_stream=img_path,
                             )
    document = part_4(document, kwargs=kwargs)
    # 6. word文档保存
    file_name = "Report of Machine_Learning_LSTM.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


if __name__ == '__main__':
    arr = np.array([1, 2, 3, 4, 5, 6])
    opt_algorithm = 1
    rul_pre = 1
    K = 1
    weights = "string"
    pso_pop_size = 5
    pso_max_itr = 5
    ga_pop_size = 5
    ga_max_itr = 5
    sa_alpha = 0.5
    sa_max_itr = 5
    accuracy = 0.99

    # word_knn(save_path="./test",
    #          train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #          opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #          K=K, weights=weights,
    #          pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #          ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #          sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #          accuracy=accuracy,
    #          )
    C = 1
    gamma = 1
    # word_svm(save_path="./test",
    #          train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #          opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #          C=C, gamma=gamma,
    #          pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #          ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #          sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #          accuracy=accuracy,
    #          )

    max_depth = 5
    max_leaf_nodes = 5
    # word_dt(save_path="./test",
    #          train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #          opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #          max_depth=max_depth, max_leaf_nodes=max_leaf_nodes,
    #          pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #          ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #          sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #          accuracy=accuracy,
    #          )

    n_estimators = 5
    # word_rf(save_path="./test",
    #         train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #         opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #         max_depth=max_depth, max_leaf_nodes=max_leaf_nodes, n_estimators=n_estimators,
    #         pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #         ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #         sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #         accuracy=accuracy,
    #         )

    # word_bagging(save_path="./test",
    #              train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #              opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #              max_depth=max_depth, max_leaf_nodes=max_leaf_nodes, n_estimators=n_estimators,
    #              pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #              ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #              sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #              accuracy=accuracy,
    #              )
    # word_et(save_path="./test",
    #         train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #         opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #         max_depth=max_depth, max_leaf_nodes=max_leaf_nodes, n_estimators=n_estimators,
    #         pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #         ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #         sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #         accuracy=accuracy,
    #         )
    LayerCount = 2
    epochs = 5
    batchSize = 128
    denseActivation = "sigmoid"

    # word_ae(save_path="./test",
    #         train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #         opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #         LayerCount=LayerCount, epochs = epochs,
    #         batchSize = batchSize, denseActivation = denseActivation,
    #         pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #         ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #         sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #         accuracy=accuracy,
    #         )

    Dropout= 0.5
    LearningRate_RBM = 0.5
    LearningRate_nn=0.5
    # word_dbn(save_path="./test",
    #          train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #          opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #          Dropout=Dropout, LearningRate_RBM=LearningRate_RBM,
    #          LearningRate_nn=LearningRate_nn,
    #          pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #          ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #          sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #          accuracy=accuracy,
    #          )

    dropout = 0.5
    batch_size = 128
    learning_rate = 0.02
    # word_cnn(save_path="./test",
    #          train_data=arr, train_label=arr, test_data=arr, test_label=arr,
    #          opt_algorithm=opt_algorithm, rul_pre=rul_pre,
    #          dropout=dropout, learning_rate=learning_rate,
    #          batch_size=batch_size, epochs=epochs,
    #          pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
    #          ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
    #          sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
    #          accuracy=accuracy,
    #          )

    lstm_count = 2
    dropoutRate = 0.5
    word_lstm(save_path="./test",
              train_data=arr, train_label=arr, test_data=arr, test_label=arr,
              opt_algorithm=opt_algorithm, rul_pre=rul_pre,
              lstm_count=lstm_count, dropoutRate=dropoutRate,
              epochs=epochs, batchSize=batchSize, denseActivation=denseActivation,
              pso_pop_size=pso_pop_size, pso_max_itr=pso_max_itr,
              ga_pop_size=ga_pop_size, ga_max_itr=ga_max_itr,
              sa_alpha=sa_alpha, sa_max_itr=sa_max_itr,
              accuracy=accuracy,
              )

