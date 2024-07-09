import time
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single"/double, "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell(cell):
    set_cell_border(
        cell,
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"}
    )


def table_full(table_rows, table):
    """
    填充表格
    :param table_rows: 表格内容
    :param table: 表格对象
    :return:
    """
    for i, row in enumerate(table_rows):
        # 获取该行所有单元格
        row_cells = table.rows[i].cells

        for g in range(len(row_cells)):
            row_cells[g].text = row[g]
            set_cell(row_cells[g])


def word_FFT(inputdata, outputdata, critical_freqs, mode, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of fast Fourier transform (FFT)", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for FFT", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Min", str(critical_freqs[0]) + " " + "Hz"),
        ("Max", str(critical_freqs[1]) + " " + "Hz"),
        ("Mode", str(mode)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "FFT.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + '.mat', "Filtered data after FFT"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + '.xlsx', "Filtered data after FFT"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + '.npy', "Filtered data after FFT"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".csv", "Filtered data after FFT"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".txt", "Filtered data after FFT"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".png", "Comparison of signals before and after FFT"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".jpg", "Comparison of signals before and after FFT"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".svg", "Comparison of signals before and after FFT"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("FFT" + ".pdf", "Comparison of signals before and after FFT"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of FFT.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_DWT_simple_filter(inputdata, outputdata, name, N, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of discrete wavelet transform (DWT) simple filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for DWT simple filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Wavelet basis", str(name)),
        ("Level", str(N)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "DWT.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_simple_filter" + '.mat', "Filtered data after DWT simple filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_simple_filter" + '.xlsx', "Filtered data after DWT simple filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_simple_filter" + '.npy', "Filtered data after DWT simple filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_simple_filter" + ".csv", "Filtered data after DWT simple filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_simple_filter" + ".txt", "Filtered data after DWT simple filter"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".png", "Comparison of signals before and after DWT simple filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".jpg", "Comparison of signals before and after DWT simple filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".svg", "Comparison of signals before and after DWT simple filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".pdf", "Comparison of signals before and after DWT simple filter"),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of DWT simple filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_DWT_Fs(inputdata, outputdata, name, F, Fs, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of discrete wavelet transform (DWT) frequency band filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for DWT frequency band filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Sample parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Sampling frequency", str(F) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("2. Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter Value"),
        ("Wavelet basis", str(name)),
        ("Max", str(Fs[1]) + " " + "Hz"),
        ("Min", str(Fs[0]) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "DWT.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_frequency_band_filter" + '.mat', "Filtered data after DWT frequency band filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_frequency_band_filter" + '.xlsx', "Filtered data after DWT frequency band filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_frequency_band_filter" + '.npy', "Filtered data after DWT frequency band filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_frequency_band_filter" + ".csv", "Filtered data after DWT frequency band filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_frequency_band_filter" + ".txt", "Filtered data after DWT frequency band filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".png", "Comparison of signals before and after DWT frequency band filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".jpg", "Comparison of signals before and after DWT frequency band filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".svg", "Comparison of signals before and after DWT frequency band filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".pdf", "Comparison of signals before and after DWT frequency band filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of DWT frequency band filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_DWT_kurtosis_max(inputdata, outputdata, name, N, k, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of discrete wavelet transform (DWT) kurtosis filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for DWT kurtosis filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Wavelet basis", str(name)),
        ("Level", str(N)),
        ("Nums of kurtosis", str(k)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "DWT.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_kurtosis_filter" + '.mat', "Filtered data after DWT kurtosis filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_kurtosis_filter" + '.xlsx', "Filtered data after DWT kurtosis filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_kurtosis_filter" + '.npy', "Filtered data after DWT kurtosis filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_kurtosis_filter" + ".csv", "Filtered data after DWT kurtosis filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_kurtosis_filter" + ".txt", "Filtered data after DWT kurtosis filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".png", "Comparison of signals before and after DWT kurtosis filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".jpg", "Comparison of signals before and after DWT kurtosis filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".svg", "Comparison of signals before and after DWT kurtosis filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".pdf", "Comparison of signals before and after DWT kurtosis filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of DWT kurtosis filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_DWT_threshold(inputdata, outputdata, name, N, threshold_method, threshold_coeff, output_file, output_image,
                       save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of discrete wavelet transform (DWT) threshold filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for DWT threshold filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    if threshold_method == 0:
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Wavelet basis", str(name)),
            ("Level", str(N)),
            ("Method", "Soft threshold"),
        ]
    elif threshold_method == 1:
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Wavelet basis", str(name)),
            ("Level", str(N)),
            ("Method", "Hard threshold"),
        ]
    else:
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Wavelet basis", str(name)),
            ("Level", str(N)),
            ("Method", "Intermediate threshold"),
            ("Threshold-coeff", str(threshold_coeff)),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "DWT.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_threshold_filter" + '.mat', "Filtered data after DWT threshold filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_threshold_filter" + '.xlsx', "Filtered data after DWT threshold filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_threshold_filter" + '.npy', "Filtered data after DWT threshold filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_threshold_filter" + ".csv", "Filtered data after DWT threshold filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT_threshold_filter" + ".txt", "Filtered data after DWT threshold filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".png", "Comparison of signals before and after DWT threshold filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".jpg", "Comparison of signals before and after DWT threshold filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".svg", "Comparison of signals before and after DWT threshold filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("DWT" + ".pdf", "Comparison of signals before and after DWT threshold filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of DWT threshold filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_EMD(inputdata, outputdata, N, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of empirical mode decomposition (EMD)", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for EMD", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Number of IMFs selected", str(N)),

    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("1. Time-domain signals", level=2)
    image_path = os.path.join(save_path, "EMD_filtered_data.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)

    document.add_heading("2. EMD decomposition", level=2)
    image_path = os.path.join(save_path, "EMD_IMFs.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_filter" + '.mat', "Filtered data after EMD"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_filter" + '.xlsx', "Filtered data after EMD"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_filter" + '.npy', "Filtered data after EMD"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_filter" + ".csv", "Filtered data after EMD"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_filter" + ".txt", "Filtered data after EMD"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_IMFs" + ".png", "IMFs after EMD decomposition"),
            ("EMD_filtered_data" + ".png", "Comparison of signals before and after EMD"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_IMFs" + ".jpg", "IMFs after EMD decomposition"),
            ("EMD_filtered_data" + ".png", "Comparison of signals before and after EMD"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_IMFs" + ".svg", "IMFs after EMD decomposition"),
            ("EMD_filtered_data" + ".png", "Comparison of signals before and after EMD"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("EMD_IMFs" + ".pdf", "IMFs after EMD decomposition"),
            ("EMD_filtered_data" + ".png", "Comparison of signals before and after EMD"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of EMD.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_mean(inputdata, outputdata, filt_length, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of mean filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for mean filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Sliding window length", str(filt_length)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "mean_filter_data.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Mean_filter" + '.mat', "Filtered data after mean filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Mean_filter" + '.xlsx', "Filtered data after mean filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Mean_filter" + '.npy', "Filtered data after mean filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("Mean_filter" + ".csv", "Filtered data after mean filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Mean_filter" + ".txt", "Filtered data after mean filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("mean_filter_data" + ".png", "Comparison of signals before and after mean filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("mean_filter_data" + ".jpg", "Comparison of signals before and after mean filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("mean_filter_data" + ".svg", "Comparison of signals before and after mean filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("mean_filter_data" + ".pdf", "Comparison of signals before and after mean filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of mean filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_WP_simple_filter(inputdata, outputdata, name, N, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of wavelet packet simple filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for wavelet packet simple filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Wavelet basis", str(name)),
        ("Level", str(N)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "WaveletPacket.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_simple_filter" + '.mat', "Filtered data after wavelet packet simple filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_simple_filter" + '.xlsx', "Filtered data after wavelet packet simple filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_simple_filter" + '.npy', "Filtered data after wavelet packet simple filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_simple_filter" + ".csv", "Filtered data after wavelet packet simple filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_simple_filter" + ".txt", "Filtered data after wavelet packet simple filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".png", "Comparison of signals before and after wavelet packet simple filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".jpg", "Comparison of signals before and after wavelet packet simple filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".svg", "Comparison of signals before and after wavelet packet simple filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".pdf", "Comparison of signals before and after wavelet packet simple filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of wavelet packet simple filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_WP_kurtosis_max(inputdata, outputdata, name, N, k, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of wavelet packet kurtosis filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for wavelet packet kurtosis filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Wavelet basis", str(name)),
        ("Level", str(N)),
        ("Nums of kurtosis", str(k)),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "WaveletPacket.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    # inline_shape=document.add_picture(image_path_or_stream=path1)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_kurtosis_filter" + '.mat', "Filtered data after wavelet packet kurtosis filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_kurtosis_filter" + '.xlsx', "Filtered data after wavelet packet kurtosis filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_kurtosis_filter" + '.npy', "Filtered data after wavelet packet kurtosis filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_kurtosis_filter" + ".csv", "Filtered data after wavelet packet kurtosis filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_kurtosis_filter" + ".txt", "Filtered data after wavelet packet kurtosis filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".png", "Comparison of signals before and after wavelet packet kurtosis filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".jpg", "Comparison of signals before and after wavelet packet kurtosis filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".svg", "Comparison of signals before and after wavelet packet kurtosis filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".pdf", "Comparison of signals before and after wavelet packet kurtosis filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of wavelet packet kurtosis filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_WP_Fs(inputdata, outputdata, name, sampling_frequency, Fs_band, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of wavelet packet frequency band filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for wavelet packet frequency band filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Sample parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Sampling frequency", str(sampling_frequency) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("2. Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Wavelet basis", str(name)),
        ("Max", str(Fs_band[1]) + " " + "Hz"),
        ("Min", str(Fs_band[0]) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "WaveletPacket.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_frequency_band_filter" + '.mat',
             "Filtered data after wavelet packet frequency band filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_frequency_band_filter" + '.xlsx',
             "Filtered data after wavelet packet frequency band filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_frequency_band_filter" + '.npy',
             "Filtered data after wavelet packet frequency band filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_frequency_band_filter" + ".csv",
             "Filtered data after wavelet packet frequency band filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Wavelet_packet_frequency_band_filter" + ".txt",
             "Filtered data after wavelet packet frequency band filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".png", "Comparison of signals before and after wavelet packet frequency band filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".jpg", "Comparison of signals before and after wavelet packet frequency band filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".svg", "Comparison of signals before and after wavelet packet frequency band filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("WaveletPacket" + ".pdf", "Comparison of signals before and after wavelet packet frequency band filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of wavelet packet frequency band filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_FK(inputdata, outputdata, nlevel, Fs, mode, order, Kurtosis_figure, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of fast kurtogram", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for fast kurtogram", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Sample parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Sampling frequency", str(Fs) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("2. Function parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Level", str(nlevel)),
        ("Mode", str(mode)),
        ("Order", str(order)),
        ("Kurtogram image", str(Kurtosis_figure)),

    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    if Kurtosis_figure == True:
        document.add_heading("1. Time-domain signals", level=2)
        image_path = os.path.join(save_path, "Fast_kurtogram.png")
        inline_shape = document.add_picture(image_path_or_stream=image_path)
        inline_shape.height = int(document.inline_shapes[0].height)
        inline_shape.width = int(document.inline_shapes[0].width)

        document.add_heading("2. Kurtogram", level=2)
        image_path = os.path.join(save_path, "Kurtosis.png")
        inline_shape = document.add_picture(image_path_or_stream=image_path)
        inline_shape.height = int(document.inline_shapes[0].height)
        inline_shape.width = int(document.inline_shapes[0].width)
    else:
        document.add_heading("Time-domain signals", level=2)
        image_path = os.path.join(save_path, "Fast_kurtogram.png")
        inline_shape = document.add_picture(image_path_or_stream=image_path)
        inline_shape.height = int(document.inline_shapes[0].height)
        inline_shape.width = int(document.inline_shapes[0].width)

    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("Fast_kurtogram" + '.mat', "Filtered data after fast kurtogram"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("Fast_kurtogram" + '.xlsx', "Filtered data after fast kurtogram"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("Fast_kurtogram" + '.npy', "Filtered data after fast kurtogram"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("Fast_kurtogram" + ".csv", "Filtered data after fast kurtogram"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("Fast_kurtogram" + ".txt", "Filtered data after fast kurtogram"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        if Kurtosis_figure == True:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".png", "Comparison of signals before and after fast kurtogram"),
                ("Kurtosis" + ".png", "Kurtogram")
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".png", "Comparison of signals before and after fast kurtogram"),
            ]
    elif output_image == 1:
        if Kurtosis_figure == True:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".jpg", "Comparison of signals before and after fast kurtogram"),
                ("Kurtosis" + ".jpg", "Kurtogram")
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".jpg", "Comparison of signals before and after fast kurtogram"),
            ]
    elif output_image == 2:
        if Kurtosis_figure == True:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".svg", "Comparison of signals before and after fast kurtogram"),
                ("Kurtosis" + ".svg", "Kurtogram")
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".svg", "Comparison of signals before and after fast kurtogram"),
            ]
    else:
        if Kurtosis_figure == True:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".pdf", "Comparison of signals before and after fast kurtogram"),
                ("Kurtosis" + ".pdf", "Kurtogram")
            ]
        else:
            table_rows = [
                ("Filename", "Description"),
                ("Fast_kurtogram" + ".pdf", "Comparison of signals before and after fast kurtogram"),
            ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of fast kurtogram.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)


def word_PSD(inputdata, outputdata, fr, n_ball, d_ball, d_pitch, alpha, frequency_band_max, factor, sideband_switch,
             sampling_frequency, cut_off_frequency, filter_method, filter_order, output_file, output_image, save_path):
    # 1. 创建一个Document对象(docx="word文件路径/None:创建新文档")
    document = Document(docx=None)

    # 2. 创建文件大标题
    document.add_heading("Report of physics-based power spectral density (PSD) threshold filter", level=0)

    # 2.1 运行日志信息
    # print(time.strftime("%Y-%m-%d", time.localtime(time.time())))
    create_datum = time.strftime("%Y-%m-%d", time.localtime(time.time()))
    create_datum = f"Create date: {create_datum}"

    document.add_paragraph(text=create_datum, style=None)

    """
    注意：
    style = None/"List Bullet"(原点列表)/"List Number"(数字列表)
    """

    # 3. 子标题一: 输入数据的基本信息（如：样本的信息，长度，大小，lable; )
    document.add_heading("1. Input and output for physics-based PSD threshold filter", level=1)

    # 4. 创建表格
    # 4.1 表格数据准备
    # 注意注意！！！！：表格中内容只支持字符串！！

    table_rows = [
        ("Data", "Data shape"),
        ("input_data", str(inputdata.shape)),
        ("output_data", str(outputdata.shape)),
    ]

    # 4.2 表格对象创建
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    # 4.3 循环数据，填入表格
    table_full(table_rows, table)

    # 子标题二：方法的信息：（方法类型，参数配置，）
    # 表格2.1
    document.add_heading("2. Method information", level=1)
    document.add_heading("1. Sample parameter", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Sampling frequency", str(sampling_frequency) + " " + "Hz"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("2. Bearing parameters", level=2)
    table_rows = [
        ("Parameter", "Parameter value"),
        ("Rotation frequency", str(fr) + " " + "Hz"),
        ("Number of rolling elements", str(n_ball)),
        ("Ball diameter", str(d_ball) + " " + "mm"),
        ("Pitch diameter", str(d_pitch) + " " + "mm"),
        ("Initial contact angle", str(alpha) + " " + "°"),
    ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    document.add_heading("3. Function parameters", level=2)
    if filter_method == 0:
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Frequency band max", str(frequency_band_max) + " " + "Hz"),
            ("Factor", str(factor)),
            ("Filter mode", "Mode" + str(sideband_switch + 1)),
            ("Cut off frequency", str(cut_off_frequency) + " " + "Hz"),
            ("Filter method", "Without low-pass filter"),
            ("Filter order", str(filter_order)),
        ]
    else:
        table_rows = [
            ("Parameter", "Parameter value"),
            ("Frequency band max", str(frequency_band_max) + " " + "Hz"),
            ("Factor", str(factor)),
            ("Filter mode", "Mode" + str(sideband_switch + 1)),
            ("Cut off frequency", str(cut_off_frequency) + " " + "Hz"),
            ("Filter method", "With low-pass filter"),
            ("Filter order", str(filter_order)),
        ]

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 子标题三：结果信息；（数据，图，表）
    document.add_heading("3. Results", level=1)

    # 3.2. 放置结果图片
    document.add_heading("Time-domain signals", level=2)
    image_path = os.path.join(save_path, "PSD_filtered_data.png")
    inline_shape = document.add_picture(image_path_or_stream=image_path)
    inline_shape.height = int(document.inline_shapes[0].height)
    inline_shape.width = int(document.inline_shapes[0].width)
    # 子标题四：结果保存路径信息；（数据，图，表）
    # 数据
    document.add_heading("4. Save files", level=1)
    document.add_heading("Result data", level=2)
    if output_file == 0:
        table_rows = [
            ("Filename", "Description"),
            ("physics_based_PSD" + '.mat', "Filtered data after physics-based PSD threshold filter"),
        ]
    elif output_file == 1:
        table_rows = [
            ("Filename", "Description"),
            ("physics_based_PSD" + '.xlsx', "Filtered data after physics-based PSD threshold filter"),
        ]
    elif output_file == 2:
        table_rows = [
            ("Filename", "Description"),
            ("physics_based_PSD" + '.npy', "Filtered data after physics-based PSD threshold filter"),
        ]
    elif output_file == 3:
        table_rows = [
            ("Filename", "Description"),
            ("physics_based_PSD" + ".csv", "Filtered data after physics-based PSD threshold filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("physics_based_PSD" + ".txt", "Filtered data after physics-based PSD threshold filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)
    # 图

    document.add_heading("Result image", level=2)
    if output_image == 0:
        table_rows = [
            ("Filename", "Description"),
            ("PSD_filtered_data" + ".png", "Comparison of signals before and after physics-based PSD threshold filter"),
        ]
    elif output_image == 1:
        table_rows = [
            ("Filename", "Description"),
            ("PSD_filtered_data" + ".jpg", "Comparison of signals before and after physics-based PSD threshold filter"),
        ]
    elif output_image == 2:
        table_rows = [
            ("Filename", "Description"),
            ("PSD_filtered_data" + ".svg", "Comparison of signals before and after physics-based PSD threshold filter"),
        ]
    else:
        table_rows = [
            ("Filename", "Description"),
            ("PSD_filtered_data" + ".pdf", "Comparison of signals before and after physics-based PSD threshold filter"),
        ]
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))

    table_full(table_rows, table)

    # 4. word文档保存
    file_name = "Report of physics-based PSD threshold filter.doc"
    save_path = os.path.join(save_path, file_name)
    document.save(save_path)
